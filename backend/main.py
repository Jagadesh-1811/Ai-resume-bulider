import os
import json
import io
from pathlib import Path
from typing import List, Dict, Optional, Any
from fastapi import FastAPI, HTTPException, Request, Body, Depends, UploadFile, File, Form
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
import google.generativeai as genai
from supabase import create_client, Client
import textstat
import jwt
import PyPDF2
import docx
from datetime import datetime, timedelta, timezone
from dotenv import load_dotenv

load_dotenv()

# Setup Supabase
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")  # Anon key for frontend
# Service role key bypasses RLS - required for backend DB operations
SUPABASE_SERVICE_KEY = os.getenv("SUPABASE_SERVICE_KEY") or os.getenv("SUPABASE_SERVICE_ROLE_KEY") or SUPABASE_KEY
try:
    supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
except Exception:
    supabase = None

import re

# Setup Google Gemini API
# Get API key from environment: GEMINI_API_KEY
# You can get your key from: https://aistudio.google.com/apikey
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    print("[WARNING] GEMINI_API_KEY not set in .env file. AI features will not work.")
else:
    genai.configure(api_key=GEMINI_API_KEY)

# Use Gemini 2.5 Flash Lite model for fast, efficient responses
GEMINI_MODEL = "gemini-2.5-flash-lite"

def call_gemini_json(system_prompt: str, user_prompt: str) -> dict:
    """Call Google Gemini API and return parsed JSON response.
    
    Uses the Gemini 2.5 Flash Lite model for fast processing.
    """
    if not GEMINI_API_KEY:
        raise Exception("GEMINI_API_KEY is not configured. Set it in .env file or environment variables.")
    
    try:
        # Create model instance (genai is already configured at module load)
        model = genai.GenerativeModel(GEMINI_MODEL)
        
        # Combine system and user prompts for Gemini
        full_prompt = f"{system_prompt}\n\n{user_prompt}"
        
        response = model.generate_content(
            full_prompt,
            generation_config=genai.types.GenerationConfig(
                temperature=0.7,
                top_p=0.95,
                max_output_tokens=2048
            )
        )
        
        if not response.text:
            raise Exception("Empty response from Gemini API")
        
        content = response.text
        
        # Extract JSON block if the model wraps it in markdown code fences
        json_match = re.search(r'\{[\s\S]*\}', content)
        if json_match:
            try:
                return json.loads(json_match.group())
            except json.JSONDecodeError:
                pass
        
        # Try parsing the entire response as JSON
        try:
            return json.loads(content)
        except json.JSONDecodeError as e:
            print(f"[GEMINI] Invalid JSON response: {str(e)}")
            raise Exception(f"Invalid JSON response from Gemini: {str(e)}")
            
    except Exception as e:
        print(f"[GEMINI] Error: {str(e)}")
        with open("ai_errors.log", "a") as f:
            f.write(f"\n[{datetime.now()}] GEMINI API ERROR: {str(e)}")
        raise Exception(f"Gemini API error: {str(e)}")

def call_gemini_text(system_prompt: str, user_prompt: str) -> str:
    """Call Google Gemini API and return plain text response.
    
    Used for content enhancement and text generation.
    """
    if not GEMINI_API_KEY:
        raise Exception("GEMINI_API_KEY is not configured. Set it in .env file or environment variables.")
    
    try:
        # Create model instance (genai is already configured at module load)
        model = genai.GenerativeModel(GEMINI_MODEL)
        
        # Combine system and user prompts for Gemini
        full_prompt = f"{system_prompt}\n\n{user_prompt}"
        
        response = model.generate_content(
            full_prompt,
            generation_config=genai.types.GenerationConfig(
                temperature=0.7,
                top_p=0.95,
                max_output_tokens=2048
            )
        )
        
        if not response.text:
            raise Exception("Empty response from Gemini API")
        
        return response.text.strip()
        
    except Exception as e:
        print(f"[GEMINI] Error: {str(e)}")
        with open("ai_errors.log", "a") as f:
            f.write(f"\n[{datetime.now()}] GEMINI API ERROR: {str(e)}")
        raise Exception(f"Gemini API error: {str(e)}")

# Setup CORS Origins
cors_origins_str = os.getenv("CORS_ORIGINS", "http://localhost:5500,http://127.0.0.1:5500,http://localhost:8000,http://127.0.0.1:8000")
CORS_ORIGINS = [origin.strip() for origin in cors_origins_str.split(",") if origin.strip() and origin.strip() != "*"]
FRONTEND_DIR = Path(__file__).resolve().parent.parent / "frontend"

app = FastAPI(title="AI Resume Builder API", version="1.0.0")

# Allow Frontend CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ORIGINS,
    allow_origin_regex=r"https?://(localhost|127\.0\.0\.1)(:\d+)?$",
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# In-memory chat history (session-based, not persisted to DB)
# Key: session_id, Value: {messages: [...], interview_idx: int}
chat_memory: Dict[str, Any] = {}
MAX_CHAT_HISTORY = 20  # Keep last 20 messages per session

# --- Authentication & JWT Validation ---
security = HTTPBearer()

# --- Auth Models ---
class AuthReq(BaseModel):
    email: str
    password: str

# --- JWT Settings ---
JWT_SECRET = os.getenv("JWT_SECRET") or "your-default-secret-if-missing"
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 60 * 24 * 7  # 7 days

class AuthUser:
    def __init__(self, user_id, email):
        self.id = user_id
        self.email = email

def create_access_token(data: dict, expires_delta: timedelta = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.now(timezone.utc) + expires_delta
    else:
        expire = datetime.now(timezone.utc) + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, JWT_SECRET, algorithm=ALGORITHM)
    return encoded_jwt

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    """
    Validates the custom JWT token passed in the Authorization header.
    Returns a mock AuthUser object.
    """
    token = credentials.credentials

    # 1) Try validating as custom app JWT first
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[ALGORITHM])
        user_id: str = payload.get("sub")
        email: str = payload.get("email")
        if user_id is None:
            raise HTTPException(status_code=401, detail="Invalid token")
        return AuthUser(user_id=user_id, email=email)
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token has expired")
    except jwt.InvalidTokenError:
        pass

    # 2) Fallback: validate Supabase JWT (Google OAuth and other Supabase sessions)
    if supabase:
        try:
            supa_user_res = supabase.auth.get_user(token)
            supa_user = getattr(supa_user_res, "user", None)
            if supa_user and getattr(supa_user, "id", None):
                return AuthUser(user_id=supa_user.id, email=getattr(supa_user, "email", ""))
        except Exception:
            pass

    raise HTTPException(status_code=401, detail="Invalid token")

# --- Pydantic Models ---

class ChatMessageReq(BaseModel):
    session_id: str
    user_id: str
    resume_id: str
    user_message: str
    resume_snapshot: Dict[str, Any] = {}

class ResumeUpdateReq(BaseModel):
    resume_data: Dict[str, Any]

class ExportDocxReq(BaseModel):
    resume_data: Dict[str, Any] = {}

class EnhanceContentReq(BaseModel):
    resume_id: str
    section_type: str
    raw_text: str

class PersonalizeReq(BaseModel):
    resume_id: str
    job_description_text: str

class OAuthCompleteReq(BaseModel):
    access_token: str

# --- DB & Core Logic Helpers ---

# Fields asked in strict order — do NOT change this order
REQUIRED_RESUME_FIELDS = [
    "name", "email", "github", "linkedin",
    "summary", "skills", "projects", "education", "experience", "achievements"
]

# These fields can be skipped if user says no / doesn't have one
OPTIONAL_FIELDS = set()

# Skip values — if a field is set to one of these the field is considered filled
SKIP_VALUES = {"none", "no", "n/a", "na", "skip", "i don't have", "i don't have one", "-"}

FIELD_QUESTIONS = {
    "name":        "What is your full name?",
    "email":       "What is your Gmail address?",
    "github":      "What is your GitHub profile link?",
    "linkedin":    "What is your LinkedIn profile link?",
    "summary":     "Give me a short 2-3 line professional summary about yourself.",
    "skills":      "List your top technical skills separated by commas.",
    "projects":    "Tell me about one of your projects: name, description, tech stack, and if you have them — a GitHub repo link and a live demo link.",
    "education":   "What is your highest education? Provide the degree, university, and graduation year.",
    "experience":  "Tell me your latest work experience: role, company, duration, and key achievement.",
    "achievements": "Share your top achievements with numbers or impact if possible."
}

def _is_filled(resume_data: Dict[str, Any], field: str) -> bool:
    value = resume_data.get(field)
    if isinstance(value, list):
        return len(value) > 0
    if isinstance(value, dict):
        return len(value.keys()) > 0
    if isinstance(value, str):
        v = value.strip()
        if not v:
            return False
        # For optional fields, treat skip-values as "filled"
        if field in OPTIONAL_FIELDS and v.lower() in SKIP_VALUES:
            return True
        return True
    return value is not None

# Fields that must contain a valid URL (not plain text)
URL_FIELDS = {"github", "linkedin"}

# Known domain fragments accepted for URL fields (user may paste without http)
URL_DOMAIN_HINTS = {
    "github":      ["github.com"],
    "linkedin":    ["linkedin.com"],
}

def is_valid_url_for_field(field: str, value: str) -> bool:
    """Return True if value looks like a valid URL for the given field."""
    if not isinstance(value, str):
        return False
    v = value.strip().lower()
    if not v:
        return False
    # For optional fields, skip-values count as valid
    if field in OPTIONAL_FIELDS and v in SKIP_VALUES:
        return True
    # Must start with http/https OR contain the expected domain
    has_protocol = v.startswith("http://") or v.startswith("https://")
    domain_hints = URL_DOMAIN_HINTS.get(field, [])
    has_domain = any(hint in v for hint in domain_hints)
    return has_protocol or has_domain

def is_valid_email(value: str) -> bool:
    """Return True if value looks like a valid email address."""
    if not isinstance(value, str):
        return False
    v = value.strip()
    if not v:
        return False
    # Must contain @ with something on both sides, and a dot after @
    at_idx = v.find("@")
    if at_idx < 1:
        return False
    domain_part = v[at_idx + 1:]
    return "." in domain_part and len(domain_part) > 3

def get_missing_required_fields(resume_data: Dict[str, Any]) -> List[str]:
    """Return fields not yet answered, in strict order."""
    return [field for field in REQUIRED_RESUME_FIELDS if not _is_filled(resume_data, field)]

def normalize_extracted_data(extracted_data: Any) -> Dict[str, Any]:
    if not extracted_data or not isinstance(extracted_data, dict):
        return {}

    normalized = dict(extracted_data)

    # Alias normalization
    if "gmail" in normalized and "email" not in normalized:
        normalized["email"] = normalized.pop("gmail")

    for key in ["skills", "projects", "education", "experience", "achievements"]:
        if key in normalized and normalized[key] is not None and not isinstance(normalized[key], list):
            normalized[key] = [normalized[key]]

    # URL validation — drop invalid values for link fields so the AI re-asks
    for field in URL_FIELDS:
        if field in normalized:
            val = normalized[field]
            if isinstance(val, str) and not is_valid_url_for_field(field, val):
                # Not a real URL — remove it so the field stays unfilled
                del normalized[field]

    # Email validation — drop invalid email values so the AI re-asks
    for email_key in ("email", "gmail"):
        if email_key in normalized:
            val = normalized[email_key]
            if isinstance(val, str) and not is_valid_email(val):
                del normalized[email_key]

    return normalized

def merge_resume_data(current_resume: Dict[str, Any], extracted_data: Dict[str, Any]) -> Dict[str, Any]:
    merged = dict(current_resume)
    for k, v in extracted_data.items():
        if isinstance(v, list):
            existing = merged.get(k, [])
            if not isinstance(existing, list):
                existing = [existing]
            for item in v:
                if item not in existing:
                    existing.append(item)
            merged[k] = existing
        elif isinstance(v, dict):
            existing = merged.get(k, {})
            if not isinstance(existing, dict):
                existing = {}
            existing.update(v)
            merged[k] = existing
        else:
            # Prevent AI from accidentally wiping out existing valid data with empty strings/nulls
            if v is None:
                continue
            if isinstance(v, str) and not v.strip():
                continue
            merged[k] = v
    return merged

def ensure_message_list(value: Any) -> List[Dict[str, str]]:
    if not isinstance(value, list):
        return []
    result: List[Dict[str, str]] = []
    for item in value:
        if isinstance(item, dict):
            msg = str(item.get("message", "")).strip()
            if msg:
                result.append({"message": msg})
        elif isinstance(item, str) and item.strip():
            result.append({"message": item.strip()})
    return result

def ensure_string_list(value: Any) -> List[str]:
    if not isinstance(value, list):
        return []
    cleaned: List[str] = []
    for item in value:
        if isinstance(item, str) and item.strip():
            cleaned.append(item.strip())
        elif isinstance(item, dict):
            # Accept occasional model outputs like {"role": "Backend Engineer"}
            for v in item.values():
                if isinstance(v, str) and v.strip():
                    cleaned.append(v.strip())
                    break
    # Deduplicate while preserving order
    seen = set()
    unique = []
    for s in cleaned:
        if s not in seen:
            seen.add(s)
            unique.append(s)
    return unique

def normalize_analysis_payload(ai_analysis: Dict[str, Any]) -> Dict[str, Any]:
    if not isinstance(ai_analysis, dict):
        ai_analysis = {}
    return {
        "keyword_density": str(ai_analysis.get("keyword_density", "N/A")),
        "impact_words": ensure_string_list(ai_analysis.get("impact_words", [])),
        "gaps_detected": bool(ai_analysis.get("gaps_detected", False)),
        "issues": ensure_message_list(ai_analysis.get("issues", [])),
        "suggestions": ensure_message_list(ai_analysis.get("suggestions", [])),
        "job_role_recommendations": ensure_string_list(ai_analysis.get("job_role_recommendations", [])),
    }

def normalize_personalize_payload(result: Dict[str, Any]) -> Dict[str, Any]:
    if not isinstance(result, dict):
        result = {}
    suggested_changes = result.get("suggested_changes", {})
    if not isinstance(suggested_changes, dict):
        suggested_changes = {}
    keyword_match_rate = result.get("keyword_match_rate", 0)
    if not isinstance(keyword_match_rate, int):
        try:
            keyword_match_rate = int(keyword_match_rate)
        except Exception:
            keyword_match_rate = 0
    keyword_match_rate = max(0, min(100, keyword_match_rate))
    return {
        "keyword_match_rate": keyword_match_rate,
        "missing_keywords": ensure_string_list(result.get("missing_keywords", [])),
        "suggested_changes": {
            "summary": str(suggested_changes.get("summary", "N/A")),
            "ai_skill_recommendations": ensure_string_list(suggested_changes.get("ai_skill_recommendations", [])),
        },
        "recommended_roles": ensure_string_list(result.get("recommended_roles", [])),
    }

IN_MEMORY_RESUMES = {}

def get_or_create_resume(resume_id: str, user_id: str):
    if resume_id not in IN_MEMORY_RESUMES:
        default_data = {
            "name": "", "title": "", "email": "", "github": "", "linkedin": "",
            "summary": "",
            "experience": [], "education": [], "skills": [], "projects": [], "achievements": []
        }
        IN_MEMORY_RESUMES[resume_id] = default_data
    return IN_MEMORY_RESUMES[resume_id]

def save_resume(resume_id: str, data: Dict[str, Any]):
    IN_MEMORY_RESUMES[resume_id] = data

# --- Endpoints ---

@app.get("/api/resume/{resume_id}")
async def get_resume(resume_id: str, user=Depends(get_current_user)):
    """Fetch structured resume JSON from DB"""
    data = get_or_create_resume(resume_id, user.id)
    return {"resume": data}

@app.put("/api/resume/{resume_id}")
async def update_resume(resume_id: str, req: ResumeUpdateReq, user=Depends(get_current_user)):
    """Save structured resume JSON to DB"""
    save_resume(resume_id, req.resume_data)
    return {"status": "success"}

@app.get("/api/config")
async def get_config():
    """Returns public configuration for the frontend to initialize Supabase."""
    return {
        "supabaseUrl": SUPABASE_URL,
        "supabaseKey": SUPABASE_KEY
    }

@app.get("/api/health")
async def health_check():
    """Simple health endpoint for frontend connectivity checks."""
    return {
        "status": "ok",
        "frontend_served": FRONTEND_DIR.exists()
    }

@app.post("/api/auth/signup")
async def signup(req: AuthReq):
    """Sign up a new user via Supabase Auth"""
    if not supabase:
        raise HTTPException(status_code=500, detail="Supabase not configured. Check .env file.")
    try:
        res = supabase.auth.sign_up({"email": req.email, "password": req.password})
        if res.user is None:
            raise Exception("Signup failed")
        
        access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
        access_token = create_access_token(
            data={"sub": res.user.id, "email": req.email}, expires_delta=access_token_expires
        )
        return {"access_token": access_token, "user": {"id": res.user.id, "email": req.email}}
    except Exception as e:
        err_msg = str(e)
        if "Email not confirmed" in err_msg:
            err_msg = "Email not confirmed. Please disable 'Confirm email' in Supabase Auth settings to enable immediate login."
        with open("ai_errors.log", "a") as f:
            f.write(f"\n[{datetime.now()}] AUTH ERROR (signup): {str(e)}")
        raise HTTPException(status_code=400, detail=err_msg)

@app.post("/api/auth/login")
async def login(req: AuthReq):
    """Log in an existing user via Supabase Auth"""
    if not supabase:
        raise HTTPException(status_code=500, detail="Supabase not configured. Check .env file.")
    try:
        res = supabase.auth.sign_in_with_password({"email": req.email, "password": req.password})
        if res.user is None:
            raise Exception("Login failed")
        
        access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
        access_token = create_access_token(
            data={"sub": res.user.id, "email": req.email}, expires_delta=access_token_expires
        )
        return {"access_token": access_token, "user": {"id": res.user.id, "email": req.email}}
    except Exception as e:
        err_msg = str(e)
        if "Email not confirmed" in err_msg:
            err_msg = "Email not confirmed. Please disable 'Confirm email' in Supabase Auth settings to enable immediate login."
        with open("ai_errors.log", "a") as f:
            f.write(f"\n[{datetime.now()}] AUTH ERROR (login): {str(e)}")
        raise HTTPException(status_code=400, detail=err_msg)

@app.get("/api/auth/validate")
async def validate_token(user=Depends(get_current_user)):
    """Validate the current JWT token. Returns 200 if valid, 401 if expired/invalid."""
    return {"valid": True, "user_id": user.id, "email": user.email}

@app.post("/api/auth/oauth/complete")
async def complete_oauth_login(req: OAuthCompleteReq):
    """Validate a Supabase OAuth access token and exchange it for the app JWT."""
    if not supabase:
        raise HTTPException(status_code=500, detail="Supabase not configured. Check .env file.")

    try:
        supa_user_res = supabase.auth.get_user(req.access_token)
        supa_user = getattr(supa_user_res, "user", None)
        if not supa_user or not getattr(supa_user, "id", None):
            raise HTTPException(status_code=401, detail="Invalid OAuth session")

        access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
        app_access_token = create_access_token(
            data={"sub": supa_user.id, "email": getattr(supa_user, "email", "")},
            expires_delta=access_token_expires,
        )
        return {
            "access_token": app_access_token,
            "user": {
                "id": supa_user.id,
                "email": getattr(supa_user, "email", ""),
                "user_metadata": getattr(supa_user, "user_metadata", {}) or {},
            },
        }
    except HTTPException:
        raise
    except Exception as e:
        with open("ai_errors.log", "a") as f:
            f.write(f"\n[{datetime.now()}] AUTH ERROR (oauth_complete): {str(e)}")
        raise HTTPException(status_code=400, detail="OAuth login could not be completed")

@app.post("/api/chat/message")
async def chat_message(req: ChatMessageReq, user=Depends(get_current_user)):
    """
    Conversational AI Resume Chat Assistant.
    Questions are asked in strict fixed order: name → email → github → linkedin →
    summary → skills → projects → education → experience → achievements.
    """
    # Ensure the user only touches their own data
    if req.user_id != user.id:
        req.user_id = user.id

    # ── Restore resume data from snapshot if backend memory was wiped ──
    current_resume = get_or_create_resume(req.resume_id, user.id)
    if req.resume_snapshot:
        filled_in_memory   = sum(1 for f in REQUIRED_RESUME_FIELDS if _is_filled(current_resume, f))
        filled_in_snapshot = sum(1 for f in REQUIRED_RESUME_FIELDS if _is_filled(req.resume_snapshot, f))
        if filled_in_snapshot > filled_in_memory:
            current_resume = merge_resume_data(current_resume, req.resume_snapshot)
            save_resume(req.resume_id, current_resume)

    # ── Session state: always start at question 0 for a brand-new session ──
    if req.session_id not in chat_memory:
        chat_memory[req.session_id] = {"messages": [], "interview_idx": 0}

    session_data    = chat_memory[req.session_id]
    # Migrate old list-based sessions gracefully
    if isinstance(session_data, list):
        chat_memory[req.session_id] = {"messages": session_data, "interview_idx": 0}
        session_data = chat_memory[req.session_id]

    session_history = session_data["messages"]
    interview_idx   = session_data.get("interview_idx", 0)

    # ── Determine current field from the ordered index ──
    # Questions always proceed in REQUIRED_RESUME_FIELDS order, one at a time.
    # The index is never influenced by old snapshot data — only by actual answers.
    if interview_idx < len(REQUIRED_RESUME_FIELDS):
        next_field = REQUIRED_RESUME_FIELDS[interview_idx]
    else:
        next_field = None
    all_done = next_field is None

    next_question_text = FIELD_QUESTIONS.get(next_field, "") if next_field else ""

    url_note = (
        "For github, linkedin: ONLY extract if the user provided a proper URL "
        "(must contain the domain e.g. github.com, linkedin.com, or start with http/https). "
        "For email/gmail: Extract the email address. It must contain an '@' symbol."
    )
    out_of_bounds_note = (
        "OUT OF BOUNDS RULE: You are strictly a Resume Builder Assistant. "
        "Do NOT answer general knowledge questions, write code, or discuss topics unrelated to the user's resume. "
        "If the user asks something off-topic, politely decline and steer them back."
    )
    anti_hallucination_note = (
        "ANTI-HALLUCINATION RULE: Extract ONLY what the user explicitly provides. "
        "Do NOT invent job roles, company names, project descriptions, or skills. "
        "If the user only gives a project name, only extract the name."
    )

    if all_done:
        system_prompt = f"""You are Skill Genius AI. All resume fields have been collected.
RULES: {out_of_bounds_note}
Return JSON: {{"extracted_data": {{}}}}"""
    else:
        field_key_instruction = {
            "name":         'Put the full name in the "name" key. Example: {{"extracted_data": {{"name": "John Doe"}}}}',
            "email":        'Put the email in the "email" key. Example: {{"extracted_data": {{"email": "john@gmail.com"}}}}',
            "github":       'Put the GitHub URL in the "github" key. Example: {{"extracted_data": {{"github": "github.com/john"}}}}',
            "linkedin":     'Put the LinkedIn URL in the "linkedin" key. Example: {{"extracted_data": {{"linkedin": "linkedin.com/in/john"}}}}',
            "summary":      'Put the summary text in the "summary" key. Example: {{"extracted_data": {{"summary": "Engineer with 3 years..."}}}}',
            "skills":       'Put skills as a JSON array in the "skills" key. NEVER use the "name" key. Example: {{"extracted_data": {{"skills": ["Python","React"]}}}}',
            "projects":     'Put project details as a JSON array in the "projects" key. Example: {{"extracted_data": {{"projects": [{{"name":"App","description":"...","tech_stack":"React"}}]}}}}',
            "education":    'Put education as a JSON array in the "education" key. Example: {{"extracted_data": {{"education": [{{"degree":"B.Tech","school":"MIT","year":"2023"}}]}}}}',
            "experience":   'Put experience as a JSON array in the "experience" key. Example: {{"extracted_data": {{"experience": [{{"title":"Engineer","company":"Google","date":"2021-2023","bullets":["Built X"]}}]}}}}',
            "achievements": 'Put achievements as a JSON array in the "achievements" key. Example: {{"extracted_data": {{"achievements": ["Won hackathon","Led team of 5"]}}}}',
        }.get(next_field, f'Extract the answer into the "{next_field}" key.')

        system_prompt = f"""You are a data extraction assistant. Extract ONLY what is asked.

QUESTION BEING ANSWERED: "{next_question_text}"
FIELD TO FILL: "{next_field}"

RULES:
- {url_note}
- {out_of_bounds_note}
- {anti_hallucination_note}
- {field_key_instruction}
- Use ONLY the key shown above. Return ONLY valid JSON, nothing else.

Return: {{"extracted_data": {{"<key>": <value>}}}} or {{"extracted_data": {{}}}} if nothing to extract."""

    # Build prompt
    resume_context = f"Current resume data: {json.dumps(current_resume)}"
    history_parts = [
        f"{'Interviewer' if m.get('role') == 'assistant' else 'Candidate'}: {m.get('content', '')}"
        for m in session_history[-MAX_CHAT_HISTORY:]
    ]
    user_prompt = f"{resume_context}\n\nConversation so far:\n{chr(10).join(history_parts)}\n\nCandidate: {req.user_message}"

    # ── Call Gemini AI (extraction only) ──
    try:
        parsed = call_gemini_json(system_prompt, user_prompt)
        extracted_data = normalize_extracted_data(parsed.get("extracted_data", {}))
    except Exception as e:
        with open("ai_errors.log", "a") as f:
            f.write(f"\n[{datetime.now()}] AI ERROR (chat_message): {str(e)}")
        extracted_data = {}

    # ── Merge extracted data into resume ──
    if extracted_data:
        current_resume = merge_resume_data(current_resume, extracted_data)
        save_resume(req.resume_id, current_resume)

    # ── Advance interview index when field is filled or user explicitly skips ──
    user_said_skip = req.user_message.strip().lower() in SKIP_VALUES
    field_filled_now = (
        next_field is not None and _is_filled(current_resume, next_field)
    )

    if next_field and (field_filled_now or user_said_skip):
        interview_idx += 1
        session_data["interview_idx"] = interview_idx

    # ── Determine next question after advance ──
    if interview_idx < len(REQUIRED_RESUME_FIELDS):
        next_field_after = REQUIRED_RESUME_FIELDS[interview_idx]
        next_question    = FIELD_QUESTIONS.get(next_field_after, "")
        all_done_now     = False
    else:
        next_field_after = None
        next_question    = ""
        all_done_now     = True

    # ── Build deterministic reply ──
    if all_done_now:
        reply = "Your resume is complete! Feel free to use the Enhance buttons, or paste a job description to get tailored suggestions."
    elif field_filled_now or user_said_skip:
        reply = f"Got it. {next_question}"
    else:
        # Nothing was extracted (off-topic or bad format) — re-ask same field
        reply = f"I didn't catch that. {next_question_text}"

    # ── Persist chat history in memory ──
    session_history.append({"role": "user",      "content": req.user_message})
    session_history.append({"role": "assistant", "content": reply})
    if len(session_history) > MAX_CHAT_HISTORY * 2:
        session_data["messages"] = session_history[-(MAX_CHAT_HISTORY * 2):]

    return {
        "reply":               reply,
        "extracted_data":      extracted_data,
        "next_step":           "complete" if all_done_now else "continue_interview",
        "missing_fields":      [f for f in REQUIRED_RESUME_FIELDS if not _is_filled(current_resume, f)],
        "required_fields_done": all_done_now,
        "next_question":       next_question,
    }

@app.post("/api/resume/enhance-content")
async def enhance_content(req: EnhanceContentReq, user=Depends(get_current_user)):
    """
    AI Content Enhancer - handles summaries, bullets, and grammar correction.
    """
    section_lower = req.section_type.lower()
    is_summary = section_lower in ["summary", "professional summary"]
    is_skills = section_lower == "skills"

    # Clean the raw text — strip leading bullet chars, extra whitespace
    import unicodedata
    raw = req.raw_text.strip()
    # Remove common bullet/list characters at line starts
    clean_lines = []
    for line in raw.splitlines():
        line = line.strip()
        line = line.lstrip('•·‣⁃▪▸-–—*>\t')
        line = line.strip()
        if line:
            clean_lines.append(line)
    clean_text = '\n'.join(clean_lines)

    if not clean_text:
        clean_text = raw

    # Get current resume to provide context
    try:
        resume = get_or_create_resume(req.resume_id, user.id)
    except Exception:
        resume = {}

    context_parts = []
    if resume.get("name"): context_parts.append(f"Candidate Name: {resume['name']}")
    if resume.get("title"): context_parts.append(f"Target Role: {resume['title']}")
    if resume.get("skills"): context_parts.append(f"Top Skills: {', '.join(resume['skills']) if isinstance(resume['skills'], list) else resume['skills']}")
    context_str = ("\n\nCandidate context (use naturally in the output):\n" + "\n".join(f"- {p}" for p in context_parts)) if context_parts else ""

    if is_summary:
        system_prompt = f"""You are an expert Resume Writer. Rewrite the provided text into a polished 2-4 sentence professional summary.
Fix all grammar and spelling. Use strong, confident language with industry-relevant keywords.{context_str}

Rules:
- Output ONLY valid JSON, no extra text.
- Use exactly this format: {{"text": "<your improved summary here>"}}"""
        user_prompt = f"Rewrite this professional summary:\n{clean_text}"
    elif is_skills:
        system_prompt = f"""You are an expert Resume Skills Advisor. Given a comma-separated list of skills, expand and improve it.
- Group and add closely related, industry-relevant skills the candidate likely has based on what's provided.
- Remove duplicates. Format each skill with proper capitalisation (e.g. "Python", "React.js", "REST APIs").
- Add at most 5 relevant extra skills not already listed.{context_str}

Rules:
- Output ONLY valid JSON, no extra text.
- Use exactly this format: {{"skills": ["Skill 1", "Skill 2", "Skill 3"]}}"""
        user_prompt = f"Improve and expand this skills list:\n{clean_text}"
    else:
        system_prompt = f"""You are an expert Resume Bullet Point Writer. Rewrite the provided text into strong, professional resume bullet points.
Fix all grammar and spelling. Use strong action verbs. Add realistic, modest metrics where none are given.{context_str}

Rules:
- Output ONLY valid JSON, no extra text.
- Each bullet must be a single, complete sentence.
- Use exactly this format: {{"bullets": ["Bullet 1 here.", "Bullet 2 here."]}}"""
        user_prompt = f"Rewrite these {req.section_type} bullet points:\n{clean_text}"

    try:
        parsed = call_gemini_json(system_prompt, user_prompt)

        if is_summary:
            # Try multiple possible keys the model might use
            text = (
                parsed.get("text") or
                parsed.get("summary") or
                parsed.get("content") or
                parsed.get("professional_summary") or
                ""
            )
            # If text is a list, join it
            if isinstance(text, list):
                text = " ".join(str(t) for t in text)
            text = str(text).strip()
            if text:
                return {"generated_content": text, "is_paragraph": True}
            # Last resort: return original text unchanged
            return {"generated_content": clean_text, "is_paragraph": True}

        elif is_skills:
            skills = (
                parsed.get("skills") or
                parsed.get("skill_list") or
                parsed.get("items") or
                []
            )
            if isinstance(skills, str):
                skills = [s.strip() for s in skills.split(',') if s.strip()]
            if isinstance(skills, list):
                flat = []
                for s in skills:
                    if isinstance(s, str) and s.strip():
                        flat.append(s.strip())
                skills = flat
            if not skills:
                skills = clean_lines if clean_lines else ["Please add skills and try enhancing again."]
            return {"generated_content": skills, "is_paragraph": False}

        else:
            # Try multiple possible keys the model might use
            bullets = (
                parsed.get("bullets") or
                parsed.get("bullet_points") or
                parsed.get("points") or
                parsed.get("items") or
                []
            )
            # If bullets came back as a single string, split by newline
            if isinstance(bullets, str):
                bullets = [line.strip().lstrip('•·-* ') for line in bullets.splitlines() if line.strip()]
            # Ensure it's a flat list of strings
            if isinstance(bullets, list):
                flat = []
                for b in bullets:
                    if isinstance(b, str) and b.strip():
                        flat.append(b.strip().lstrip('•·-* '))
                    elif isinstance(b, dict):
                        val = next((v for v in b.values() if isinstance(v, str) and v.strip()), None)
                        if val: flat.append(val.strip())
                bullets = flat
            if not bullets:
                # Fallback: use cleaned original lines as bullets
                bullets = clean_lines if clean_lines else ["Please add content and try enhancing again."]
            return {"generated_content": bullets, "is_paragraph": False}

    except Exception as e:
        with open("ai_errors.log", "a") as f:
            f.write(f"\n[{datetime.now()}] AI ERROR (enhance_content): {str(e)}")
        if is_summary:
            return {"generated_content": clean_text or "Failed to enhance. Please try again.", "is_paragraph": True}
        return {"generated_content": clean_lines or ["Failed to enhance content. Please try again."], "is_paragraph": False}

def calculate_ats_metrics(resume_text: str) -> dict:
    """Calculate ATS-specific metrics for resume quality assessment."""
    if not resume_text or len(resume_text.strip()) < 50:
        return {
            "total_length": 0,
            "word_count": 0,
            "line_count": 0,
            "has_contact_info": False,
            "has_summary": False,
            "has_experience": False,
            "has_skills": False,
            "has_education": False,
            "formatting_score": 0
        }

    lines = resume_text.split('\n')
    words = resume_text.split()
    
    # Check for common resume sections
    text_lower = resume_text.lower()
    has_summary = any(keyword in text_lower for keyword in ['summary', 'professional', 'objective', 'about'])
    has_experience = any(keyword in text_lower for keyword in ['experience', 'work', 'employment'])
    has_skills = any(keyword in text_lower for keyword in ['skills', 'technical', 'expertise', 'proficiency'])
    has_education = any(keyword in text_lower for keyword in ['education', 'degree', 'university', 'college'])
    
    # Check for contact info (email and phone patterns)
    has_email = bool(re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', resume_text))
    has_phone = bool(re.search(r'[\d\-\(\) ]{10,}', resume_text))
    has_contact = has_email or has_phone

    # Formatting score (prefer plain text, avoid special formatting)
    special_chars_ratio = len(re.findall(r'[^a-zA-Z0-9\s\-\.\,]', resume_text)) / max(len(resume_text), 1)
    formatting_score = max(0, 100 - int(special_chars_ratio * 200))

    # Readability score
    readability = textstat.flesch_reading_ease(resume_text)
    
    return {
        "total_length": len(resume_text),
        "word_count": len(words),
        "line_count": len(lines),
        "has_contact_info": has_contact,
        "has_summary": has_summary,
        "has_experience": has_experience,
        "has_skills": has_skills,
        "has_education": has_education,
        "formatting_score": min(100, formatting_score),
        "readability_score": int(readability)
    }

def extract_action_verbs(resume_text: str) -> list:
    """Extract common action verbs used in resume accomplishments."""
    action_verbs = [
        "achieved", "built", "created", "designed", "developed", "directed", "driven",
        "executed", "expanded", "focused", "founded", "generated", "guided", "improved",
        "increased", "initiated", "innovated", "launched", "led", "leveraged", "managed",
        "mentored", "optimized", "organized", "pioneered", "planned", "produced", "proposed",
        "reduced", "resolved", "restructured", "redesigned", "scaled", "solved", "spearheaded",
        "strategized", "streamlined", "strengthened", "transformed", "accelerated", "established"
    ]
    
    text_lower = resume_text.lower()
    found_verbs = []
    
    for verb in action_verbs:
        pattern = r'\b' + verb + r'\b'
        if re.search(pattern, text_lower):
            found_verbs.append(verb.capitalize())
    
    return found_verbs[:10]  # Return top 10

@app.post("/api/analysis/upload")
async def upload_analysis_report(
    file: UploadFile = File(...),
    target_role: str = Form(...),
    user=Depends(get_current_user)
):
    """
    AI Resume Analysis & ATS Scoring Dashboard
    Analyzes uploaded PDF/DOCX and returns comprehensive ATS scoring against the target role.
    """
    # 1. Extract Text from Uploaded File
    resume_text = ""
    file_bytes = await file.read()
    
    try:
        if file.filename.endswith(".pdf"):
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            for page in pdf_reader.pages:
                resume_text += page.extract_text() + "\n"
        elif file.filename.endswith(".docx"):
            doc = docx.Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs:
                resume_text += para.text + "\n"
        elif file.filename.endswith(".txt"):
            resume_text = file_bytes.decode("utf-8")
        else:
            raise HTTPException(400, "Unsupported file format. Please upload PDF, DOCX, or TXT.")
    except Exception as e:
        raise HTTPException(400, f"Error reading file: {str(e)}")
        
    resume_text = resume_text.strip()
    
    # 2. Calculate ATS Metrics
    ats_metrics = calculate_ats_metrics(resume_text)
    action_verbs = extract_action_verbs(resume_text)
    
    # 3. AI Deep Analysis
    system_prompt = f"""You are an expert ATS (Applicant Tracking System) Resume Analyzer.
Analyze resumes for their ATS compatibility and alignment with job roles.
The target role is: '{target_role}'

You MUST return a JSON object with EXACTLY these keys:
1. "keyword_density": String value "Low"/"Optimal"/"High" - indicates keyword match with {target_role}
2. "ats_compatibility": Integer 0-100 - how well will this resume pass through ATS systems
3. "impact_score": Integer 0-100 - strength of accomplishment statements and impact words
4. "section_completeness": Integer 0-100 - presence of all major resume sections
5. "format_issues": Array of strings listing formatting problems (if any)
6. "missing_keywords": Array of 3-5 critical keywords missing for {target_role}
7. "issues": Array of objects {{"message": "..."}} - critical issues
8. "recommendations": Array of objects {{"message": "..."}} - specific actionable improvements
9. "strengths": Array of strings - positive aspects of the resume
10. "alternative_roles": Array of 3-5 similar job roles that fit the profile
"""
    
    try:
        if len(resume_text) > 50:
            ai_analysis = call_gemini_json(system_prompt, resume_text)
        else:
            raise Exception("Resume text is too short or empty")
    except Exception as e:
        with open("ai_errors.log", "a") as f:
            f.write(f"\n[{datetime.now()}] AI ERROR (analysis_upload): {str(e)}")
        ai_analysis = {
            "keyword_density": "N/A",
            "ats_compatibility": 50,
            "impact_score": 50,
            "section_completeness": 50,
            "format_issues": [],
            "missing_keywords": ["Please try again"],
            "issues": [{"message": "Analysis failed. Please ensure the Gemini API is configured correctly with a valid API key."}],
            "recommendations": [],
            "strengths": [],
            "alternative_roles": []
        }

    # Ensure proper normalization
    if not isinstance(ai_analysis, dict):
        ai_analysis = {}
    
    # Extract and validate numeric scores
    ats_compat = int(ai_analysis.get("ats_compatibility", 50)) if isinstance(ai_analysis.get("ats_compatibility"), (int, float)) else 50
    impact_score = int(ai_analysis.get("impact_score", 50)) if isinstance(ai_analysis.get("impact_score"), (int, float)) else 50
    section_complete = int(ai_analysis.get("section_completeness", 50)) if isinstance(ai_analysis.get("section_completeness"), (int, float)) else 50
    
    # Calculate overall score (weighted average)
    overall_score = int((ats_compat * 0.4 + impact_score * 0.3 + section_complete * 0.3))
    overall_score = max(0, min(100, overall_score))
    
    # Normalize lists
    format_issues = ensure_string_list(ai_analysis.get("format_issues", []))
    missing_keywords = ensure_string_list(ai_analysis.get("missing_keywords", []))
    issues = ensure_message_list(ai_analysis.get("issues", []))
    recommendations = ensure_message_list(ai_analysis.get("recommendations", []))
    strengths = ensure_string_list(ai_analysis.get("strengths", []))
    alt_roles = ensure_string_list(ai_analysis.get("alternative_roles", []))

    # 4. Compile Comprehensive Report
    report = {
        "overall_ats_score": overall_score,
        "detailed_scores": {
            "ats_compatibility": min(100, max(0, ats_compat)),
            "impact_and_achievement": min(100, max(0, impact_score)),
            "section_completeness": min(100, max(0, section_complete)),
            "formatting_compliance": ats_metrics.get("formatting_score", 0),
            "readability": ats_metrics.get("readability_score", 60)
        },
        "resume_structure": {
            "has_summary": ats_metrics.get("has_summary", False),
            "has_experience": ats_metrics.get("has_experience", False),
            "has_skills": ats_metrics.get("has_skills", False),
            "has_education": ats_metrics.get("has_education", False),
            "has_contact_info": ats_metrics.get("has_contact_info", False),
            "word_count": ats_metrics.get("word_count", 0)
        },
        "keyword_analysis": {
            "keyword_density": str(ai_analysis.get("keyword_density", "N/A")),
            "impact_words_found": action_verbs,
            "missing_critical_keywords": missing_keywords[:5]
        },
        "issues": issues[:10],
        "recommendations": recommendations[:10],
        "strengths": strengths[:5],
        "alternative_job_roles": alt_roles[:5],
        "format_issues": format_issues
    }
            
    return report

@app.post("/api/resume/personalize")
async def personalize_resume(req: PersonalizeReq, user=Depends(get_current_user)):
    """
    AI Job Description Analyzer & Resume Personalizer
    """
    # 1. Fetch Real Resume Data
    resume_data = get_or_create_resume(req.resume_id, user.id)
    resume_text = json.dumps(resume_data)
    
    system_prompt = """You are an ATS Job Matching Engine.
Compare the candidate's Resume JSON against the provided Job Description.
Evaluate the match and recommend actionable ways to personalize the resume for this SPECIFIC job.

You MUST return a JSON object with:
1. "keyword_match_rate": Integer 0-100 representing the approximate overlap.
2. "missing_keywords": Array of up to 5 critical hard skills/keywords from the Job Description missing from the resume.
3. "suggested_changes": Object containing:
   - "summary": A new, highly tailored Professional Summary rewriting their current one to fit this job.
   - "ai_skill_recommendations": Array of strings suggesting how they can naturally add the missing skills.
4. "recommended_roles": Array of up to 5 job roles that best match the user's profile.

Example:
{
  "keyword_match_rate": 85,
  "missing_keywords": ["Docker", "AWS"],
  "suggested_changes": {
    "summary": "Backend engineer specializing in highly scalable Python architectures...",
    "ai_skill_recommendations": ["Add 'AWSDocker' to your skills list if you have experience with it."]
    },
    "recommended_roles": ["Backend Engineer", "Python Developer"]
}"""
    
    user_prompt = f"Resume JSON:\n{resume_text}\n\nTarget Job Description:\n{req.job_description_text}"
    
    try:
        result = call_gemini_json(system_prompt, user_prompt)
    except Exception as e:
        with open("ai_errors.log", "a") as f:
            f.write(f"\n[{datetime.now()}] AI ERROR (personalize_resume): {str(e)}")
        result = {
            "keyword_match_rate": 0,
            "missing_keywords": ["Error analyzing job match. Please try again."],
            "suggested_changes": {"summary": "N/A", "ai_skill_recommendations": []},
            "recommended_roles": []
        }

    return normalize_personalize_payload(result)

@app.post("/api/resume/export-docx")
async def export_docx(req: ExportDocxReq, user=Depends(get_current_user)):
    """Generate and return a .docx Word file from resume data sent in the request body."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed. Run: pip install python-docx")

    resume = req.resume_data
    doc = Document()

    # ------- Styles -------
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ------- Header: Name -------
    name = resume.get('name') or 'Your Name'
    heading = doc.add_heading(name, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x16, 0x23, 0x2A)

    title = resume.get('title', '')
    if title:
        t = doc.add_paragraph(title)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t.runs[0].font.size = Pt(13)
        t.runs[0].font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    # ------- Contact Line -------
    contacts = []
    if resume.get('email'): contacts.append(f"Email: {resume['email']}")
    if resume.get('github'): contacts.append(f"GitHub: {resume['github']}")
    if resume.get('linkedin'): contacts.append(f"LinkedIn: {resume['linkedin']}")
    if contacts:
        cp = doc.add_paragraph(' | '.join(contacts))
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].font.size = Pt(9)
        cp.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()  # spacer

    def section_heading(title):
        p = doc.add_heading(title, level=2)
        p.runs[0].font.color.rgb = RGBColor(0x16, 0x23, 0x2A)
        return p

    # ------- Summary -------
    summary = resume.get('summary', '')
    if summary:
        section_heading('Professional Summary')
        doc.add_paragraph(summary)

    # ------- Experience -------
    experience = resume.get('experience', [])
    if experience:
        section_heading('Experience')
        for exp in experience:
            if isinstance(exp, str):
                doc.add_paragraph(exp, style='List Bullet')
            elif isinstance(exp, dict):
                role = exp.get('title') or exp.get('role', 'Role')
                company = exp.get('company', '')
                date = exp.get('date') or exp.get('duration', '')
                header = f"{role}"
                if company: header += f" — {company}"
                if date: header += f" ({date})"
                p = doc.add_paragraph()
                run = p.add_run(header)
                run.bold = True
                run.font.size = Pt(11)
                for bullet in (exp.get('bullets') or exp.get('impact') or []):
                    doc.add_paragraph(str(bullet), style='List Bullet')

    # ------- Skills -------
    skills = resume.get('skills', [])
    if skills:
        section_heading('Skills')
        doc.add_paragraph(', '.join(str(s) for s in skills))

    # ------- Education -------
    education = resume.get('education', [])
    if education:
        section_heading('Education')
        for edu in education:
            if isinstance(edu, str):
                doc.add_paragraph(edu, style='List Bullet')
            elif isinstance(edu, dict):
                degree = edu.get('degree') or edu.get('title', 'Degree')
                school = edu.get('school') or edu.get('university', '')
                year = edu.get('year') or edu.get('date', '')
                p = doc.add_paragraph()
                run = p.add_run(degree)
                run.bold = True
                if school:
                    p.add_run(f" — {school}")
                if year:
                    p.add_run(f" ({year})")

    # ------- Projects -------
    projects = resume.get('projects', [])
    if projects:
        section_heading('Projects')
        for proj in projects:
            if isinstance(proj, str):
                doc.add_paragraph(proj, style='List Bullet')
            elif isinstance(proj, dict):
                proj_name = proj.get('name', 'Project')
                proj_desc = proj.get('description') or proj.get('summary', '')
                demo_link = proj.get('demo_link') or proj.get('demo', '')
                gh_link   = proj.get('github_link') or proj.get('github', '')
                p = doc.add_paragraph()
                run = p.add_run(proj_name)
                run.bold = True
                if proj_desc:
                    doc.add_paragraph(proj_desc)
                links = []
                if demo_link: links.append(f"Demo: {demo_link}")
                if gh_link:   links.append(f"Code: {gh_link}")
                if links:
                    lp = doc.add_paragraph(" | ".join(links))
                    lp.runs[0].font.size = Pt(9)
                    lp.runs[0].font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    # ------- Achievements -------
    achievements = resume.get('achievements', [])
    if achievements:
        section_heading('Achievements')
        for ach in achievements:
            text = ach if isinstance(ach, str) else (ach.get('title') or ach.get('description', ''))
            if text:
                doc.add_paragraph(text, style='List Bullet')

    # ------- Stream as download -------
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_name = (resume.get('name') or 'resume').replace(' ', '_')
    filename = f"{safe_name}_resume.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )

if FRONTEND_DIR.exists():
    app.mount("/", StaticFiles(directory=FRONTEND_DIR, html=True), name="frontend")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
